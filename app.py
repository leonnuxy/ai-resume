import os
from flask import Flask, render_template, request, redirect, url_for, flash, session, send_file, jsonify
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import re
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.utils import simpleSplit
import json
from typing import Dict, List, Tuple
from datetime import timedelta
from flask_session import Session
import pdfplumber
import ollama

app = Flask(__name__)
app.secret_key = os.urandom(24)  # Generate a strong random key
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)  # Increase session lifetime
app.config['SESSION_TYPE'] = 'filesystem'  # Use filesystem to store session data
Session(app)  # Initialize session

# Configure upload folder
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Create uploads directory if it doesn't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

class ResumeSection:
    def __init__(self, name: str, y_position: float, font_size: float):
        self.name = name
        self.y_position = y_position
        self.font_size = font_size
        self.content: List[Tuple[str, float]] = []  # (text, font_size)
    
    def to_dict(self) -> Dict:
        """Convert ResumeSection to dictionary for JSON serialization"""
        return {
            'name': self.name,
            'y_position': self.y_position,
            'font_size': self.font_size,
            'content': self.content
        }
    
    @classmethod
    def from_dict(cls, data: Dict) -> 'ResumeSection':
        """Create ResumeSection from dictionary"""
        section = cls(data['name'], data['y_position'], data['font_size'])
        section.content = data['content']
        return section

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def analyze_document_structure(text: str) -> dict:
    """Use Ollama to analyze document structure and formatting"""
    try:
        prompt = f"""Analyze this document and provide formatting recommendations in JSON format.
        Focus on:
        1. Identifying section headers
        2. Determining relative font sizes for different elements
        3. Suggesting spacing between sections
        4. Identifying key formatting patterns
        
        Document text:
        {text}
        
        Respond only with JSON in this format:
        {{
            "sections": [
                {{
                    "text": "section text",
                    "is_header": boolean,
                    "relative_size": float,
                    "spacing_after": float
                }}
            ],
            "base_font_size": float,
            "formatting_notes": string
        }}
        """
        
        response = ollama.chat(model='mistral', messages=[{
            'role': 'user',
            'content': prompt
        }])
        
        # Extract JSON from response
        analysis = json.loads(response['message']['content'])
        return analysis
    except Exception as e:
        print(f"AI analysis error: {str(e)}")
        return {}

def process_word(word: dict, base_size: float = None) -> dict:
    """Process a single word and return its formatting."""
    if not word['text'].strip():
        return {}
    
    size = word['size']
    normalized_size = round((size / base_size) * 10, 1) if base_size else 10
    
    return {
        'text': word['text'].strip(),
        'font_size': normalized_size,
        'top': word['top'],
        'bottom': word['bottom']
    }

def calculate_line_spacing(sections: list) -> dict:
    """Calculate line spacing between sections."""
    spacing = {}
    sorted_sections = sorted(sections, key=lambda x: x['top'])
    
    for i in range(len(sorted_sections)-1):
        curr, next_sect = sorted_sections[i], sorted_sections[i+1]
        space = next_sect['top'] - curr['bottom']
        if space > 0:
            spacing[curr['text']] = min(space / 2, 10)
    
    return spacing

def extract_pdf_formatting(file_path: str) -> dict:
    formatting = {
        'sections': [],
        'default_font_size': 10,
        'line_spacing': {},
        'ai_analysis': None
    }
    
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                words = page.extract_words(keep_blank_chars=True)
                text += page.extract_text() + "\n"
                
                sizes = [word['size'] for word in words if word['text'].strip()]
                base_size = sorted(set(sizes), key=sizes.count)[-1] if sizes else None
                
                for word in words:
                    section = process_word(word, base_size)
                    if section:
                        formatting['sections'].append(section)
            
            formatting['ai_analysis'] = analyze_document_structure(text)
            formatting['line_spacing'] = calculate_line_spacing(formatting['sections'])
            
    except Exception as e:
        print(f"Error extracting formatting: {str(e)}")
    
    return formatting

def parse_pdf(file_path):
    """Parse PDF and extract both text and formatting"""
    text = ""
    formatting = extract_pdf_formatting(file_path)
    
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    
    return text, formatting

def parse_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def parse_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


@app.before_request
def make_session_permanent():
    session.permanent = True

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    session.permanent = True
    
    if 'resume' not in request.files:
        flash('No file selected')
        return redirect(url_for('index'))
    
    file = request.files['resume']
    if file.filename == '':
        flash('No file selected')
        return redirect(url_for('index'))
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        try:
            file_ext = filename.rsplit('.', 1)[1].lower()
            formatting = None
            
            if file_ext == 'pdf':
                text, formatting = parse_pdf(file_path)
            elif file_ext == 'docx':
                text = parse_docx(file_path)
            elif file_ext == 'txt':
                text = parse_txt(file_path)
            
            # Anonymize personal information in text
            modified_text = text
            name_pattern = r'\b[A-Z][a-z]+ [A-Z][a-z]+\b'
            modified_text = re.sub(name_pattern, 'Test Name', modified_text)
            
            # Clear and update session data
            session.clear()
            session['modified_text'] = modified_text
            session['original_file_type'] = file_ext
            session['original_filename'] = filename
            if formatting:
                session['formatting'] = formatting
            
            os.remove(file_path)
            
            return render_template('result.html', content=modified_text)
            
        except Exception as e:
            print(f"Error processing file: {str(e)}")
            flash('Error processing file')
            return redirect(url_for('index'))
    else:
        flash('Invalid file type. Please upload a PDF, DOCX, or TXT file.')
        return redirect(url_for('index'))

def create_txt_file(text: str) -> Tuple[io.BytesIO, str]:
    """Create a text file in memory."""
    return io.BytesIO(text.encode('utf-8')), 'text/plain'

def create_docx_file(text: str) -> Tuple[io.BytesIO, str]:
    """Create a Word document in memory."""
    doc = Document()
    doc.add_paragraph(text)
    file_data = io.BytesIO()
    doc.save(file_data)
    file_data.seek(0)
    return file_data, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

def apply_text_formatting(text: str, format_info: dict) -> tuple:
    """Apply formatting to text and return font settings."""
    font_size = format_info.get('size', 10)
    spacing = format_info.get('spacing', 2)
    is_header = format_info.get('is_header', False)
    font_name = "Helvetica-Bold" if is_header else "Helvetica"
    return font_name, font_size, spacing

def get_text_format(line: str, text_format_map: dict, default_font_size: int = 10) -> dict:
    """Get formatting for a line of text."""
    for original_text, format_info in text_format_map.items():
        if line.strip() in original_text or original_text in line.strip():
            return format_info
    return {'size': default_font_size, 'spacing': 2, 'is_header': False}

def draw_text_line(canvas, line: str, y: float, margin: int, font_name: str, font_size: float) -> float:
    """Draw a line of text and return new y position."""
    if not line.strip():
        return y - font_size
    
    canvas.setFont(font_name, font_size)
    canvas.drawString(margin, y, line)
    return y

def process_formatted_text(c, text: str, text_format_map: dict, y: float, margin: int, default_font_size: int) -> float:
    """Process text with formatting and return new y position."""
    for line in text.split('\n'):
        if y < 50:
            c.showPage()
            y = 750
        
        format_info = get_text_format(line, text_format_map, default_font_size)
        font_name, font_size, spacing = apply_text_formatting(line, format_info)
        y = draw_text_line(c, line, y, margin, font_name, font_size)
        y -= (font_size + spacing)
    return y

def process_simple_text(c, text: str, y: float, margin: int, default_font_size: int) -> float:
    """Process text without formatting and return new y position."""
    for line in text.split('\n'):
        if y < 50:
            c.showPage()
            y = 750
            c.setFont("Helvetica", default_font_size)
        
        y = draw_text_line(c, line, y, margin, "Helvetica", default_font_size)
        y -= (default_font_size + 2)
    return y

def create_formatted_pdf(text: str, formatting: dict = None):
    file_data = io.BytesIO()
    c = canvas.Canvas(file_data, pagesize=letter)
    y = 750
    margin = 72
    default_font_size = 10
    
    if formatting and 'sections' in formatting:
        ai_analysis = formatting.get('ai_analysis', {})
        ai_sections = {section['text']: section for section in ai_analysis.get('sections', [])} if ai_analysis else {}
        
        text_format_map = {
            section['text']: {
                'size': section['font_size'] / 2 * (1.2 if ai_sections.get(section['text'], {}).get('is_header', False) else 1),
                'spacing': formatting['line_spacing'].get(section['text'], 2),
                'is_header': ai_sections.get(section['text'], {}).get('is_header', False)
            }
            for section in formatting['sections']
        }
        
        for paragraph in text.split('\n\n'):
            y = process_formatted_text(c, paragraph, text_format_map, y, margin, default_font_size)
            y -= 10
    else:
        c.setFont("Helvetica", default_font_size)
        for paragraph in text.split('\n\n'):
            y = process_simple_text(c, paragraph, y, margin, default_font_size)
            y -= 10
    
    c.save()
    file_data.seek(0)
    return file_data, 'application/pdf'

@app.route('/check_session')
def check_session():
    return '', 204

@app.route('/download')
def download_file():
    try:
        if 'modified_text' not in session:
            flash('No processed file available for download')
            return redirect(url_for('index'))

        modified_text = session['modified_text']
        formatting = session.get('formatting', None)
        
        # Create formatted PDF with the actual modified text and formatting
        file_data, mimetype = create_formatted_pdf(modified_text, formatting)
        
        # Use original filename if available, otherwise default to alt_resume.pdf
        original_filename = session.get('original_filename', 'alt_resume.pdf')
        download_name = f"modified_{original_filename}"
        if not download_name.endswith('.pdf'):
            download_name = download_name.rsplit('.', 1)[0] + '.pdf'

        return send_file(
            file_data,
            mimetype=mimetype,
            as_attachment=True,
            download_name=download_name
        )
    except Exception as e:
        print(f"Error in download: {str(e)}")
        flash('Error generating download file')
        return redirect(url_for('index'))

def analyze_resume_for_job(resume_text: str, job_description: str) -> dict:
    """Use Ollama to analyze and optimize resume for a specific job"""
    try:
        prompt = f"""You are a professional resume optimization expert. Your goal is to analyze the provided resume and job description and provide specific, actionable suggestions to make the resume a stronger match for the job.
            
            Analyze the following RESUME and JOB DESCRIPTION to provide resume optimization suggestions.
            Focus your analysis on these key areas:

            1. Identify Missing Key Skills/Experiences: Determine crucial skills or experiences listed in the JOB DESCRIPTION that are absent or underemphasized in the RESUME. Suggest how the candidate can address these gaps.
            2. Suggest Targeted Improvements for Job Matching:  Recommend specific rewrites or additions to the RESUME to better align it with the JOB DESCRIPTION's requirements and preferences.
            3. Highlight and Emphasize Relevant Strengths: Identify existing experiences, skills, or accomplishments in the RESUME that are highly relevant to the JOB DESCRIPTION but might be currently understated. Advise on how to emphasize these points for maximum impact.
            4. Provide General Formatting and Wording Improvements: Offer any overall suggestions for enhancing the resume's clarity, readability, formatting, and professional tone to make it more appealing to recruiters for this specific job.

            Please provide your response in a structured JSON format, detailing your suggestions within these categories:
            
            Resume:
            {resume_text}
            
            Job Description:
            {job_description}
            
            Respond with specific, actionable suggestions in JSON format:
            {{
                "missing_skills": [
                    {{
                        "skill": "Specific skill or experience missing from resume",
                        "suggestion": "Actionable advice on how to incorporate or address this gap in the resume (e.g., add to skills section, rephrase experience to highlight this skill, etc.)"
                    }}
                ],
                "improvement_suggestions": [
                    {{
                        "current": "Extract of the current text or resume section that needs improvement",
                        "suggested": "Improved or rewritten version of the text/section",
                        "reason": "Brief explanation of why this change strengthens the resume's alignment with the job description"
                    }}
                ],
                "emphasis_suggestions": [
                    {{
                        "experience": "Specific experience or accomplishment from the resume",
                        "why_relevant": "Explanation of direct relevance to the job description's requirements",
                        "how_to_emphasize": "Concrete suggestion on how to better highlight this experience (e.g., move to top of section, use stronger action verbs, quantify achievements, etc.)"
                    }}
                ],
                "general_suggestions": [
                    "Overall, general advice for improving the resume's effectiveness for this job application (e.g., check for keywords, tailor summary, improve formatting of skills section, etc.)"
                ]
            }}
        """
        
        response = ollama.chat(model='mistral', messages=[{
            'role': 'user',
            'content': prompt
        }])
        
        # Extract JSON from response
        analysis = json.loads(response['message']['content'])
        return analysis
    except Exception as e:
        print(f"AI analysis error: {str(e)}")
        return {}

def format_optimization_suggestions(analysis: dict) -> str:
    """Format the AI analysis into HTML"""
    if not analysis:
        return "<p>Unable to generate optimization suggestions.</p>"
    
    CLOSE_UL = "</ul>"
    html = []
    
    # Missing Skills
    if analysis.get('missing_skills'):
        html.append("<h4>Missing Key Skills</h4><ul>")
        for skill in analysis['missing_skills']:
            html.append(f"<li><strong>{skill['skill']}</strong>: {skill['suggestion']}</li>")
        html.append(CLOSE_UL)
    
    # Improvement Suggestions
    if analysis.get('improvement_suggestions'):
        html.append("<h4>Suggested Improvements</h4><ul>")
        for suggestion in analysis['improvement_suggestions']:
            html.append(f"<li><div class='highlight'>{suggestion['current']}</div> could be improved to: ")
            html.append(f"<div class='highlight'>{suggestion['suggested']}</div>")
            html.append(f"<em>Why: {suggestion['reason']}</em></li>")
        html.append(CLOSE_UL)
    
    # Emphasis Suggestions
    if analysis.get('emphasis_suggestions'):
        html.append("<h4>Experiences to Emphasize</h4><ul>")
        for emphasis in analysis['emphasis_suggestions']:
            html.append(f"<li><strong>{emphasis['experience']}</strong>")
            html.append(f"<br>Relevance: {emphasis['why_relevant']}")
            html.append(f"<br>Suggestion: {emphasis['how_to_emphasize']}</li>")
        html.append(CLOSE_UL)
    
    # General Suggestions
    if analysis.get('general_suggestions'):
        html.append("<h4>General Suggestions</h4><ul>")
        for suggestion in analysis['general_suggestions']:
            html.append(f"<li>{suggestion}</li>")
        html.append(CLOSE_UL)
    
    return "\n".join(html)

@app.route('/optimize', methods=['POST'])
def optimize_resume():
    try:
        if 'modified_text' not in session:
            flash('No resume available for optimization')
            return redirect(url_for('index'))
        
        job_description = request.form.get('job_description')
        if not job_description:
            flash('Please provide a job description')
            return redirect(url_for('result'))
        
        # Get the resume text from session
        resume_text = session['modified_text']
        
        # Analyze resume against job description
        analysis = analyze_resume_for_job(resume_text, job_description)
        
        # Format the suggestions
        optimized_content = format_optimization_suggestions(analysis)
        
        # Store the job description and analysis in session
        session['job_description'] = job_description
        session['optimization_analysis'] = analysis
        
        return render_template('result.html',
                            content=resume_text,
                            optimized_content=optimized_content,
                            session_data=dict(session))
                            
    except Exception as e:
        print(f"Error in optimization: {str(e)}")
        flash('Error during resume optimization')
        return redirect(url_for('index'))

if __name__ == '__main__':
    try:
        response = ollama.chat(model='mistral', messages=[{
            'role': 'user',
            'content': 'Simple test message'
        }])
    except Exception:
        pass
    
    app.run(debug=False) 