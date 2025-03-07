import pdfplumber
from docx import Document
import re
import os  # Import os for path manipulation
from typing import Tuple, Dict, List  # Import Tuple, Dict, and List


def process_word(word: dict, base_size: float = None) -> dict:
    """Process a single word and return its formatting."""
    if not isinstance(word, dict) or 'text' not in word or not word['text'].strip() or 'size' not in word:
        return {}

    size = word['size']
    normalized_size = round((size / base_size) * 10, 1) if base_size else 10

    return {
        'text': word['text'].strip(),
        'font_size': normalized_size,
        'top': word['top'],
        'bottom': word['bottom']
    }

def calculate_line_spacing(sections: list) -> dict:
    """Calculate line spacing between sections."""
    spacing = {}
    sorted_sections = sorted(sections, key=lambda x: x['top'])

    for i in range(len(sorted_sections) - 1):
        curr, next_sect = sorted_sections[i], sorted_sections[i + 1]
        space = next_sect['top'] - curr['bottom']
        if space > 0:
            spacing[curr['text']] = min(space / 2, 10)

    return spacing

def extract_pdf_formatting(file_path: str) -> dict:
    """Extract formatting information from a PDF file."""
    formatting = {
        'sections': [],
        'default_font_size': 10,
        'line_spacing': {},
        'ai_analysis': None,
        'section_analysis': None,
        'text': ""  # Initialize text here
    }
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                try:  # Added a try-except block here
                    words = page.extract_words(keep_blank_chars=True)
                    formatting['text'] += page.extract_text() + "\n"

                    sizes = [word['size'] for word in words if isinstance(word, dict) and 'size' in word and word['text'].strip()]
                    base_size = sorted(set(sizes), key=sizes.count, reverse=True)[0] if sizes else None

                    for word in words:
                        section = process_word(word, base_size)
                        if section:
                            formatting['sections'].append(section)

                except Exception as e: # Inner Exception
                    print(f"Error processing page {page.page_number}: {str(e)}")
                    continue # Continue with next page

            # Get AI analysis (moved outside the page loop)
            # Check if ai_analysis module and functions are available before calling
            if 'ai_analysis' in formatting:
                try:
                    from app.services.ai_analysis import analyze_document_structure, analyze_resume_sections
                    formatting['ai_analysis'] = analyze_document_structure(formatting['text'])
                    formatting['section_analysis'] = analyze_resume_sections(formatting['text'])

                    # Apply AI section analysis (moved outside the page loop)
                    if formatting['section_analysis'] and 'sections' in formatting['section_analysis']:
                        for section in formatting['section_analysis']['sections']:
                            for fmt_section in formatting['sections']:
                                if section['name'].lower() in fmt_section.get('text', '').lower():
                                    fmt_section['is_header'] = section.get('formatting', {}).get('is_header', False)
                                    fmt_section['font_size'] = max(fmt_section.get('font_size', 10),
                                                                section.get('formatting', {}).get('suggested_font_size', 10))
                except ImportError:
                    print("AI analysis module not found. Skipping AI analysis.")
                    pass  # Continue without AI analysis


    except Exception as e: # Outer Exception
        print(f"Error extracting formatting: {str(e)}")

    return formatting

def parse_pdf(file_path: str) -> Tuple[str, dict]:
    """Parse PDF and extract both text and formatting"""
    formatting = extract_pdf_formatting(file_path) #get complete formatting
    text = formatting.get('text', "") # Get the text extracted by pdfplumber

    return text, formatting

def parse_docx(file_path: str) -> Tuple[str, dict]:
    """Parse DOCX file and extract text and basic formatting."""
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    text = '\n'.join(full_text)

    # Create a simplified formatting dictionary, similar to the PDF version
    formatting = {
        'sections': [],
        'default_font_size': 11,  # Common default font size for DOCX
        'line_spacing': {}, # DOCX line spacing can be complex, skipping for now
        'ai_analysis': None,
        'section_analysis': None,
        'text': text
    }

    # Create "sections" based on paragraphs (very basic)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # Only add non-empty paragraphs
             formatting['sections'].append({
                'text': paragraph.text.strip(),
                'font_size': formatting['default_font_size'],  # Use default
                'top': 0,  # No positional info in this simple extraction
                'bottom': 0,
            })

    return text, formatting

def parse_txt(file_path: str) -> Tuple[str, dict]:
    """Parse TXT file and extract text."""
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    # Create a simplified formatting dictionary, similar to the PDF version
    formatting = {
        'sections': [],
        'default_font_size': 10,
        'line_spacing': {},
        'ai_analysis': None,
        'section_analysis': None,
        'text': text
    }
    return text, formatting

def parse_resume(file_path: str) -> Tuple[str, dict]:
    """Parse resume based on file extension, return text and formatting."""
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == '.pdf':
        return parse_pdf(file_path)
    elif file_extension == '.docx':
        return parse_docx(file_path)
    elif file_extension == '.txt':
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {file_extension}")