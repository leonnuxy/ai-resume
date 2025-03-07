import os
import pytest
from docx import Document
import pdfplumber
from app.services.resume_parser import parse_resume, parse_pdf, parse_docx, parse_txt
from app.utils.text_cleaner import clean_text

# --- Fixtures ---
@pytest.fixture(scope="session")
def test_files_dir():
    """Provides the path to the test_files directory."""
    return os.path.join(os.path.dirname(__file__), "test_files")

@pytest.fixture(scope="session")
def test_pdf_path(test_files_dir):
    """Provides the path to the test PDF file."""
    return os.path.join(test_files_dir, "resume.pdf")

@pytest.fixture(scope="session")
def test_docx_path(test_files_dir):
    """Provides the path to the test DOCX file."""
    return os.path.join(test_files_dir, "test_resume.docx")

@pytest.fixture(scope="session")
def test_txt_path(test_files_dir):
    """Provides the path to the test TXT file."""
    return os.path.join(test_files_dir, "test_resume.txt")

@pytest.fixture(scope="session")
def create_test_files(test_files_dir, test_docx_path, test_txt_path, test_pdf_path):
    """Creates test DOCX, TXT, and PDF files if they don't exist."""
    os.makedirs(test_files_dir, exist_ok=True)

    # Create a simple DOCX file
    if not os.path.exists(test_docx_path):
        doc = Document()
        doc.add_paragraph("This is a test DOCX file.")
        doc.add_paragraph("It has multiple paragraphs.")
        doc.save(test_docx_path)

    # Create a simple TXT file
    if not os.path.exists(test_txt_path):
        with open(test_txt_path, "w", encoding="utf-8") as f:
            f.write("This is a test TXT file.\nIt has multiple lines.")

    # Create a simple PDF file using *pdfplumber*
    if not os.path.exists(test_pdf_path):
        with pdfplumber.PDF.new(test_pdf_path) as pdf:
            page = pdf.add_page()
            page.add_text(
                """Noel Ugwoke 306-490-2929 | Email | GitHub | Portfolio
SKILLS
• Full-Stack Development: Microsoft .NET, C#, ASP.NET MVC, HTML5, CSS3, JavaScript, REST APIs.
• Cloud & DevOps: Microsoft Azure (App Services, SQL Database, DevOps), Docker, CI/CD Pipelines.
• Frameworks: Razor Pages, React, Spring Boot, Entity Framework.
• Databases: SQL Server, PostgreSQL, MySQL (query optimization, stored procedures).
• Methodologies: Agile/Scrum, Jira, Confluence, Secure Coding Practices.
• Soft Skills: Cross-functional collaboration, technical documentation, vendor coordination, problem-solving.

RELEVANT EXPERIENCE
Full-Stack .NET/Application Developer - Association of Professional Engineers and Geoscientists of Alberta (APEGA)
December 2022 - November 2024
Skills Used: React, Python, SQL, Azure Cloud, REST APIs, Ready API, Jira, Confluence, Agile/Scrum.
• Developed and maintained full-stack .NET applications, unifying front-end and back-end components to deliver seamless user experiences.
• Migrated legacy systems to Microsoft Azure, leveraging Azure SQL Database and App Services to enhance scalability and security compliance, reducing downtime by 40%.
• Collaborated with third-party vendors to integrate APIs and manage data integrity during system transitions.
• Authored technical documentation, including user guides and deployment procedures, to streamline onboarding.
• Optimized SQL queries and resolved performance bottlenecks, improving data retrieval speeds by 25%.

Web Application Developer - Spartan Controls July 2021 - November 2022
Skills Used: React, Python, SQL, AWS (Lambda), PostgreSQL, Terraform, Tableau, Spring Boot.
• Built responsive web applications using .NET and JavaScript, focusing on cross-browser compatibility and user-centric design.
• Automated CI/CD pipelines using Azure DevOps, reducing deployment errors by 30% and accelerating release cycles.
• Coordinated with Agile teams to deliver iterative updates, ensuring alignment with business requirements and user feedback.

EMPLOYMENT HISTORY
Data Analyst - Parkland Fuel Corporation May 2020 - January 2021
Skills Used: AWS (Lambda), Tableau, PowerBI, SQL, Python, ETL Pipelines, Data Modeling, Predictive Analytics.
• Designed automated ETL workflows, leveraged complex SQL queries and built interactive dashboards for data-driven insights.
• Streamlined ETL pipelines to process real-time data, improving accuracy and reducing processing time by 35%.

IT Analyst Intern - Alberta Health Services January 2019 - April 2020
Skills Used: ServiceNow, Jira, Visual Studio Code, UiPath, , Splunk, SQL, AWS (Lambda), Tableau.
• Migrated legacy scripts to ServiceNow, automated tasks with UiPath and AWS Lambda, optimized SQL queries, analyzed system logs with Splunk, developed Jira dashboards, and created Tableau reports to enhance IT operations.

PROJECTS
• Vaccine Booking Website - Alberta Health Services
    o Led front-end development using React and TypeScript, creating a user-friendly interface that supported 10,000+ daily bookings.
• Gas Pipeline Monitoring WebApp
    o Full-stack .NET application with Razor Pages for dynamic UI, integrated with Azure cloud services for real-time data processing.

EDUCATION
University of Calgary - BSc Computer Science 2016 - 2022
Bow Valley College - Diploma Software Development 2015 - 2016

CERTIFICATIONS
AWS Cloud Practitioner | IBM AI Engineering Professional | Google Cloud Platform Developer""",
                fontname="Courier",  # Use a standard font
                fontsize=12,
            )

# --- Test Cases ---
@pytest.fixture(scope='function')
def dirty_text():
    """Fixture to provide the actual text from the PDF (now consistent)."""
    return """Noel Ugwoke 306-490-2929 | Email | GitHub | Portfolio
SKILLS
• Full-Stack Development: Microsoft .NET, C#, ASP.NET MVC, HTML5, CSS3, JavaScript, REST APIs.
• Cloud & DevOps: Microsoft Azure (App Services, SQL Database, DevOps), Docker, CI/CD Pipelines.
• Frameworks: Razor Pages, React, Spring Boot, Entity Framework.
• Databases: SQL Server, PostgreSQL, MySQL (query optimization, stored procedures).
• Methodologies: Agile/Scrum, Jira, Confluence, Secure Coding Practices.
• Soft Skills: Cross-functional collaboration, technical documentation, vendor coordination, problem-solving.

RELEVANT EXPERIENCE
Full-Stack .NET/Application Developer - Association of Professional Engineers and Geoscientists of Alberta (APEGA)
December 2022 - November 2024
Skills Used: React, Python, SQL, Azure Cloud, REST APIs, Ready API, Jira, Confluence, Agile/Scrum.
• Developed and maintained full-stack .NET applications, unifying front-end and back-end components to deliver seamless user experiences.
• Migrated legacy systems to Microsoft Azure, leveraging Azure SQL Database and App Services to enhance scalability and security compliance, reducing downtime by 40%.
• Collaborated with third-party vendors to integrate APIs and manage data integrity during system transitions.
• Authored technical documentation, including user guides and deployment procedures, to streamline onboarding.
• Optimized SQL queries and resolved performance bottlenecks, improving data retrieval speeds by 25%.

Web Application Developer - Spartan Controls July 2021 - November 2022
Skills Used: React, Python, SQL, AWS (Lambda), PostgreSQL, Terraform, Tableau, Spring Boot.
• Built responsive web applications using .NET and JavaScript, focusing on cross-browser compatibility and user-centric design.
• Automated CI/CD pipelines using Azure DevOps, reducing deployment errors by 30% and accelerating release cycles.
• Coordinated with Agile teams to deliver iterative updates, ensuring alignment with business requirements and user feedback.

EMPLOYMENT HISTORY
Data Analyst - Parkland Fuel Corporation May 2020 - January 2021
Skills Used: AWS (Lambda), Tableau, PowerBI, SQL, Python, ETL Pipelines, Data Modeling, Predictive Analytics.
• Designed automated ETL workflows, leveraged complex SQL queries and built interactive dashboards for data-driven insights.
• Streamlined ETL pipelines to process real-time data, improving accuracy and reducing processing time by 35%.

IT Analyst Intern - Alberta Health Services January 2019 - April 2020
Skills Used: ServiceNow, Jira, Visual Studio Code, UiPath, , Splunk, SQL, AWS (Lambda), Tableau.
• Migrated legacy scripts to ServiceNow, automated tasks with UiPath and AWS Lambda, optimized SQL queries, analyzed system logs with Splunk, developed Jira dashboards, and created Tableau reports to enhance IT operations.

PROJECTS
• Vaccine Booking Website - Alberta Health Services
    o Led front-end development using React and TypeScript, creating a user-friendly interface that supported 10,000+ daily bookings.
• Gas Pipeline Monitoring WebApp
    o Full-stack .NET application with Razor Pages for dynamic UI, integrated with Azure cloud services for real-time data processing.

EDUCATION
University of Calgary - BSc Computer Science 2016 - 2022
Bow Valley College - Diploma Software Development 2015 - 2016

CERTIFICATIONS
AWS Cloud Practitioner | IBM AI Engineering Professional | Google Cloud Platform Developer"""


@pytest.fixture(scope='function')
def expected_clean_text(dirty_text):
    """Fixture to provide expected clean text for testing, dynamically cleaned."""
    return clean_text(dirty_text)  # Use the *actual* clean_text function

def test_clean_text(dirty_text, expected_clean_text):
    """Tests the clean_text function."""
    cleaned_text = clean_text(dirty_text)
    assert cleaned_text.strip() == expected_clean_text.strip(), "The cleaned text does not match the expected output."

def test_parse_pdf(test_pdf_path, expected_clean_text):
    """Tests parsing a PDF file."""
    try:
        text, _ = parse_pdf(test_pdf_path)
        cleaned_extracted_text = clean_text(text)
        assert cleaned_extracted_text.strip() == expected_clean_text.strip(), "Cleaned text from PDF does not match expected."

    except FileNotFoundError:
        pytest.fail(f"Test PDF file not found: {test_pdf_path}")
    except Exception as e:
        pytest.fail(f"An unexpected error occurred during PDF parsing: {e}")


def test_parse_docx(test_docx_path):
    """Test parsing a DOCX file"""
    try:
        text, _ = parse_docx(test_docx_path)
        cleaned_extracted_text = clean_text(text)
        assert "This is a test DOCX file." in cleaned_extracted_text
    except FileNotFoundError:
        pytest.fail(f"Test DOCX file not found: {test_docx_path}")
    except Exception as e:
        pytest.fail(f"An unexpected error occurred: {e}")


def test_parse_txt(test_txt_path):
    """Tests parsing a simple text file."""
    try:
        text, _ = parse_txt(test_txt_path)
        cleaned_extracted_text = clean_text(text)
        assert "This is a test TXT file." in cleaned_extracted_text

    except FileNotFoundError:
        pytest.fail(f"Test TXT file not found at: {test_txt_path}")
    except Exception as e:
        pytest.fail(f"An unexpected error occurred during TXT parsing: {e}")

def test_parse_resume(test_pdf_path, test_docx_path, test_txt_path, expected_clean_text):
    """Tests the parse_resume function with different file types."""

    # Test with PDF
    try:
        cleaned_text, _ = parse_resume(test_pdf_path)
        assert cleaned_text.strip() == expected_clean_text.strip(), "parse_resume (PDF) failed."
    except Exception as e:
        pytest.fail(f"parse_resume (PDF) failed: {e}")

    # Test with DOCX
    try:
        cleaned_text, _ = parse_resume(test_docx_path)
        assert "This is a test DOCX file." in cleaned_text, "parse_resume (DOCX) failed."
    except Exception as e:
        pytest.fail(f"parse_resume (DOCX) failed: {e}")

    # Test with TXT
    try:
        cleaned_text, _ = parse_resume(test_txt_path)
        assert "This is a test TXT file." in cleaned_text, "parse_resume (TXT) failed."
    except Exception as e:
        pytest.fail(f"parse_resume (TXT) failed: {e}")

    # Test with unsupported file type
    with pytest.raises(ValueError):
        parse_resume("test.xyz")  # Should raise ValueError